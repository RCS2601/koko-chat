from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_documentation():
    doc = Document()

    # Style configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Ultimate Store & Koko Chat Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary Section
    doc.add_heading('1. Project Summary', level=1)
    summary_text = (
        "Ultimate Store is a modern, AI-powered e-commerce marketplace that redefines the shopping "
        "experience through conversation. At its core is \"Koko\", a cheeky and context-aware AI shopping "
        "assistant powered by Groq (Llama 3.3). Unlike traditional e-commerce sites where users browse "
        "grids of products, Ultimate Store allows users to chat with Koko to find products, get recommendations, "
        "and place orders in natural language."
    )
    doc.add_paragraph(summary_text)

    # Image Section
    doc.add_heading('2. Interface Preview', level=1)
    doc.add_paragraph("Below is a visual representation of the AI Chat Interface:")
    
    image_path = "/home/stoy/.gemini/antigravity/brain/0474dd59-c9a8-4844-b25a-202af8352027/koko_chat_ui_mockup_1770660818037.png"
    if os.path.exists(image_path):
        try:
            doc.add_picture(image_path, width=Inches(5.0))
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Error adding image: {str(e)}]")
    else:
        doc.add_paragraph("[Image not found]")
    
    doc.add_paragraph("Figure 1: The mobile chat interface showing Koko suggesting products.", style='Caption')

    # How It Works Section
    doc.add_heading('3. How It Works', level=1)
    
    doc.add_heading('3.1 The Tech Stack', level=2)
    doc.add_paragraph(
        "The project is built using a modern but lightweight stack:"
    )
    p = doc.add_paragraph()
    p.add_run("• Frontend: ").bold = True
    p.add_run("Vanilla JavaScript, HTML5, CSS3 (No heavy frameworks like React/Vue).")
    
    p = doc.add_paragraph()
    p.add_run("• Backend/Database: ").bold = True
    p.add_run("Firebase Firestore for real-time data synchronization.")
    
    p = doc.add_paragraph()
    p.add_run("• AI Engine: ").bold = True
    p.add_run("Groq API running Llama 3.3 70B for extremely fast inference.")

    doc.add_heading('3.2 The AI Brain (Koko)', level=2)
    ai_text = (
        "The application doesn't just send user text to an LLM. It employs a sophisticated 'System Prompt' "
        "strategy. When the app loads, the entire product catalog is fetched from Firebase and injected into "
        "the AI's context. This allows Koko to:\n"
        "1. Know exactly what products are available in real-time.\n"
        "2. Recommend products based on price, category, or description.\n"
        "3. Detect user intent (e.g., 'buy', 'search', 'chat').\n"
        "4. Respond with structured JSON data to trigger UI actions (like showing a product card)."
    )
    doc.add_paragraph(ai_text)

    doc.add_heading('3.3 King Mode', level=2)
    king_text = (
        "A unique feature called 'King Mode' (Alpha) changes the AI's persona to that of a loyal butler. "
        "It also enables 'One-Click Auto Checkout', streamlining the buying process for VIP users by "
        "bypassing confirmation screens and using default payment settings."
    )
    doc.add_paragraph(king_text)

    # Features Section
    doc.add_heading('4. Key Features', level=1)
    
    doc.add_heading('User Side (Buyer)', level=2)
    buyer_features = [
        "Natural Language Search: 'I need a cheap laptop for gaming'.",
        "Interactive Cart: Add items directly from the chat.",
        "Order Tracking: Real-time status updates.",
        "Direct Chat: Communicate with sellers about specific orders."
    ]
    for feature in buyer_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Admin Side (Seller)', level=2)
    seller_features = [
        "Dashboard Overview: Track orders, revenue, and product counts.",
        "Product Management: Add/Edit/Delete products real-time.",
        "Order Fulfillment: Confirm or reject orders showing up instantly.",
        "Availability Toggle: Mark items as out of stock with one click."
    ]
    for feature in seller_features:
        doc.add_paragraph(feature, style='List Bullet')

    # Save
    output_path = "/home/stoy/Documents/chattobott/Project_Documentation.docx"
    doc.save(output_path)
    print(f"Documentation saved to: {output_path}")

if __name__ == "__main__":
    create_documentation()
